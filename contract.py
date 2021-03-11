import docx
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

# In this section we define the contract. We put %s for the terms we need to fill in
title = "LEGAL SERVICES AGREEMENT\n\n"
parties = "IDENTIFICATION OF PARTIES. This agreement, is made between {}, (Attorney) and {}, (Client)."
services = "LEGAL SERVICES TO BE PROVIDED. The legal services to be provided by attorney to client are as follows:\n{}"
responsibilities = "RESPONSIBILITIES OF ATTORNEY AND CLIENT. Attorney will perform the legal services called for " \
    "under this agreement, keep the Client informed of progress and developments, and respond promptly to Client's " \
    "inquiries and communications. Client will be truthful and cooperative to Attorney, keep Attorney reasonably well " \
    "informed of developments and of Client's address, telephone number and whereabouts; and timely make any payments " \
    "required by this agreement."

compensation = "COMPENSATION. Client will pay the Attorney for the legal services provided under this agreement as " \
    "follows: \nHourly Compensation. In consideration for the services to be performed by Attorney, Client agrees " \
    "to pay to Attorney at the following rate: ${} per hour for legal services"
compensation2 = "Attorney will charge in increments of one tenth on an hour, rounded off for each particular activity "\
    "to the nearest one tenth of an hour. The minimum time charged for any particular activity will be one tenth of " \
    "an hour."
compensation3 = "Attorney will charge for all activities undertaken in providing legal services to Client under this " \
    "agreement, including, but not limited to, the following: conferences, court sessions, and depositions preparation " \
    "and participation; correspondence and legal documents review and preparation; legal research; and telephone " \
    "conversations. When two or more of Attorney's personnel are engaged in working on the matter at the same time, "\
    "such as in conferences between them, the time of each will be charged at his or her hourly rate."
compensation4 = "Payment is expected for all services and expenses upon the receipt of any invoice.\n\nClient " \
    "acknowledge that Attorney has made no promises about the total sum of Attorney's fees to be incurred by Client " \
    "under this agreement."
costs = "COSTS. Client will pay all \"costs\" in connection with Attorney's representation of Client under this " \
    "agreement. Costs will be advanced by Attorney and then billed to Client unless the costs can be met out of Client " \
    "deposits that are applicable towards costs. Costs include, but are not limited to, court filing fees, deposition " \
    "costs, expert frees and expenses, investigation costs, long distance telephone charges, messenger service fees, " \
    "photocopying expenses, and process server fees."
deposit = "DEPOSIT. Client will pay to Attorney and initial deposit of ${}, to be received by Attorney on or before {} "\
    "and to be applied against attorney's fees and costs incurred by Client. Of this amount ${} is refundable and ${} " \
    "is nonrefundable. The nonrefundable portion will be applied against attorney's fees first. If, at the termination " \
    "of services under this agreement, the total amount incurred by Client for attorney's fees is less than the amount " \
    "of the initial deposit, the difference, to a maximum of the refundable portion of the deposit, will be refunded " \
    "to Client."
provisions = "GENERAL PROVISIONS. This agreement sets forth the entire understanding of the parties. Any amendments " \
    "must be in writing and signed by both parties. This agreement shall be construed under the laws of {}. If any " \
    "provision of this agreement is held to be invalid, illegal or unenforceable, the remaining portions of this " \
    "agreement shall remain in full force and effect and construed so as to best effectuate the original intent " \
    "and purpose of this agreement."
effective_date = "EFFECTIVE DATE OF AGREEMENT. This agreement becomes effective as of the date it is executed by the " \
    "parties to do so.\n"
foregoing = "The foregoing is agreed to by:\n"
signatures = "___________________________________________  ______________________\n" \
             "Client Signature                                                              Date\n\n" \
             "___________________________________________  ______________________\n" \
             "Attorney Signature                                                          Date\n"

# ---------------------------------------------------------------------------------------------------------------

# in this section we give values to our items
lawyer = "Law firm of Valentina, John, Andres and Eduardo and Partners"
client = "Hana Montana"
service = "Represent Client in her music career."
compensation_value = "10000"
deposit_value = "15000"
refundable_deposit_value = "10000"
nonrefundable_deposit_value = "5000"
deposit_date = "March 11, 2025"
jurisdiction = "United States of America"

# ---------------------------------------------------------------------------------------------------------------

# start generating the document
doc = docx.Document()
# title
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = para.add_run(title)
run.font.name = "Times New Roman"
run.font.size = Pt(20)
run.font.bold = True

# parties
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = para.add_run(parties.format(lawyer, client))
run.font.name = "Times New Roman"
run.font.size = Pt(12)

# service
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = para.add_run(services.format(service))
run.font.name = "Times New Roman"
run.font.size = Pt(12)

# responsibilities
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = para.add_run(responsibilities)
run.font.name = "Times New Roman"
run.font.size = Pt(12)

# compensation
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = para.add_run(compensation.format(compensation_value))
run.font.name = "Times New Roman"
run.font.size = Pt(12)

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = para.add_run(compensation2)
run.font.name = "Times New Roman"
run.font.size = Pt(12)

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = para.add_run(compensation3)
run.font.name = "Times New Roman"
run.font.size = Pt(12)

para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = para.add_run(compensation4)
run.font.name = "Times New Roman"
run.font.size = Pt(12)

# costs
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = para.add_run(costs)
run.font.name = "Times New Roman"
run.font.size = Pt(12)

# deposit
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = para.add_run(deposit.format(deposit_value, deposit_date, refundable_deposit_value, nonrefundable_deposit_value))
run.font.name = "Times New Roman"
run.font.size = Pt(12)

# provisions
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = para.add_run(provisions.format(jurisdiction))
run.font.name = "Times New Roman"
run.font.size = Pt(12)

# effective date
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = para.add_run(effective_date)
run.font.name = "Times New Roman"
run.font.size = Pt(12)

# foregoing
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = para.add_run(foregoing)
run.font.name = "Times New Roman"
run.font.size = Pt(12)

# signatures
para = doc.add_paragraph()
para.alignment = WD_ALIGN_PARAGRAPH.LEFT
run = para.add_run(signatures)
run.font.name = "Times New Roman"
run.font.size = Pt(12)

# save the doc
doc.save(f"{client}.docx")